# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 标题 =====
title = doc.add_heading('新官网应用中心 — 内容确认清单', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('时间：2026年3月__日   地点：____   参会：产品推广国内、产品推广国际、数字智能部', style='Normal')
doc.add_paragraph('')

# ===== 〇、应用总览页 =====
doc.add_heading('〇、应用总览页样式确认（application-center.html）', level=1)

doc.add_paragraph('总览页为6个应用领域的入口页面，结构如下：', style='Normal')
doc.add_paragraph('')

doc.add_heading('页面结构', level=2)

t0a = doc.add_table(rows=7, cols=3, style='Table Grid')
t0a.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['区域', '内容', '确认点']):
    cell = t0a.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d0a = [
    ['面包屑', '首页 > 应用中心', ''],
    ['标题区', 'h1「应用中心」+ 一段引言描述', '引言文案是否OK？'],
    ['6大应用领域卡片', '3列网格(grid-cols-3)，每卡片含：Material图标 + 领域名 + 简介 + 特性标签 + 推荐系列', '每张卡片的简介、标签、系列名称是否需要修改？'],
    ['热门终端应用', '4个快捷入口卡片：安全气囊ECU / ABS防抱死 / ADAS自动驾驶 / GPU加速卡VRM', '① 安全气囊链接到 application-airbag.html（已删除！需确认）\n② ABS/ADAS 链接带 ?tab= 参数跳转\n③ 入口数量和内容是否合适？'],
    ['技术资源横幅', '深蓝底白色字 CTA — 「AI服务器PDN优化指南」下载白皮书', '① 只展示AI服务器白皮书是否合适？\n② 下载链接待补充'],
    ['设计工具入口', '3个横排卡片：寿命推算 / SPICE模型 / 3D-CAD', '链接均为 # 占位'],
]
for i, row_data in enumerate(d0a):
    for j, val in enumerate(row_data):
        t0a.rows[i+1].cells[j].text = val

doc.add_paragraph('')
doc.add_heading('6大领域卡片内容核对', level=2)

t0b = doc.add_table(rows=7, cols=4, style='Table Grid')
t0b.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['领域', '当前简介', '当前标签', '当前系列']):
    cell = t0b.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d0b = [
    ['汽车电子', '主驱/电驱电控、充电系统、安全部件、ADAS、热管理、智能座舱、车灯、充电桩 —— AEC-Q200合规，耐高温135℃，抗振动。', 'AEC-Q200 / 135℃耐高温 / 极低漏电流', '8个应用子领域 · 10+产品系列'],
    ['AI服务器', 'GPU 加速卡、CPU 供电、PDN 总线架构的大电流、超低 ESR 电容方案，支持 400V/800V DC-Link。', '大电流 / 超低ESR / 小型化', 'VHT · NPX · CW3H系列'],
    ['仪器仪表', '精密测量仪器、工业自动化仪表、医疗检测设备等对精度和稳定性要求极高的电容方案。', '高精度 / 低漏电流 / 长寿命', 'VKM · VPT · LK系列'],
    ['新型电机驱动', '变频器、伺服驱动器、步进电机驱动等工业电机控制系统的 DC-Link 与滤波电容方案。', 'DC-Link / 滤波 / 高纹波', 'VPG · VPX · CW3H系列'],
    ['储能', '光伏逆变器、储能系统、充电桩等高压大容量电容方案，支持高纹波电流与长寿命需求。', '高压大容量 / 高纹波 / 长寿命', 'CW3H · VPG系列'],
    ['消费类电子', '笔记本电脑、智能家居、PD快充、LED照明等消费电子产品的薄型化小型化电容方案。', '薄型化 / 小型化 / 低ESR', 'NPX · VKM · VPT系列'],
]
for i, row_data in enumerate(d0b):
    for j, val in enumerate(row_data):
        t0b.rows[i+1].cells[j].text = val

doc.add_paragraph('')
doc.add_heading('⚠️ 总览页需修复的问题', level=2)

issues = [
    '1. application-airbag.html 已删除 —「热门终端应用」中"安全气囊ECU"卡片链接到不存在的页面，需移除或替换为有效链接',
    '2. ?tab= 参数跳转 — ABS/ADAS 卡片通过URL参数跳转到汽车电子页指定Tab，需确认汽车电子页支持此参数（目前 automotive 页有 getUrlParam 函数，但 automotive 的 TAB_MAP 中 ADAS 对应 key 为 "adas"，需要验证）',
    '3. 白皮书横幅 — 下载按钮无实际链接，需提供PDF地址',
    '4. 设计工具3个入口 — 链接均为 #，需补充实际URL',
    '5. 6张领域卡片的数据是硬编码的 — 如果后续修改了领域简介/标签/系列名，需要同步修改总览页',
]
for item in issues:
    doc.add_paragraph(item, style='Normal')

doc.add_paragraph('')

# ===== 一、样式确认 =====
doc.add_heading('一、应用领域详情页样式确认（6页统一模板）', level=1)

t1 = doc.add_table(rows=8, cols=3, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['#', '确认项', '说明']
for i, h in enumerate(headers):
    cell = t1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data1 = [
    ['1', '整体布局', '面包屑 → 概述(标题+标签+下载按钮) → 一级Tab栏 → 左侧(二级子应用+设计工具) → 右侧内容 → 技术文章'],
    ['2', '样式1：推荐系列', '«系统模块详解与推荐系列» — 2列卡片(grid-cols-2)，每卡片含：模块图标+名称+简介+推荐系列标签'],
    ['3', '样式2：推荐规格', '«系统模块详解与推荐规格» — 每个模块占一整行，上方模块介绍，下方紧跟规格表(系列/料号/电压/容量/尺寸/ESR/纹波/寿命/备注)'],
    ['4', '电路拓扑图区', '当前6页全部使用 anquanqinang.png 占位图，标注"演示用，正式版替换"'],
    ['5', '设计工具侧边栏', '寿命推算 / SPICE模型 / 3D-CAD 三个快捷入口，链接待补充'],
    ['6', '底部技术文章', '每页5-6篇文章标题，链接均为 # 占位'],
    ['7', '一级Tab切换', 't0/t1... 键 + TAB_MAP映射，点击切换子应用列表和右侧内容'],
]
for i, row_data in enumerate(data1):
    for j, val in enumerate(row_data):
        t1.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# ===== 二、各页面数据现状 =====
doc.add_heading('二、各页面数据现状（真实 vs 占位）', level=1)

doc.add_heading('✅ 已有真实规格数据（SPECS覆盖率 100%）', level=2)

t2 = doc.add_table(rows=5, cols=6, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['页面', '一级Tab数', '子应用总数', 'SPECS规格', '模块卡片(DB)', '备注']
for i, h in enumerate(h2):
    cell = t2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d2 = [
    ['AI服务器与数据中心', '3', '7', '7/7 全部有料号', '全部有', '唯一完全覆盖的页面'],
    ['储能', '3', '4', '4/4 全部有料号', '全部有', ''],
    ['新型电机驱动', '6', '6', '6/6 全部有料号', '全部有', ''],
    ['仪器仪表', '3', '9', '9/9 全部有料号', '全部有', '含智能电表7个子应用'],
]
for i, row_data in enumerate(d2):
    for j, val in enumerate(row_data):
        t2.rows[i+1].cells[j].text = val

doc.add_paragraph('')
doc.add_heading('⚠️ 部分覆盖 / 需补充', level=2)

t3 = doc.add_table(rows=3, cols=5, style='Table Grid')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
h3 = ['页面', '子应用总数', '有SPECS', '缺口', '问题']
for i, h in enumerate(h3):
    cell = t3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d3 = [
    ['汽车电子', '36', '~11个', '25个子应用无规格数据', 't3(车灯)/t4(智能驾驶)/t5(多媒体)/t6(辅助)/t7(OBC)/t8(充电桩) 共6个Tab完全无SPECS数据；部分子应用名称与数据库键名不一致'],
    ['消费类电子', '6', '0', '6个子应用均无法匹配', 'appData子应用名(如"快充输入端专用")与SPECS键名(如"PD快充-高压输入端")不一致；电子笔/数码3C/安防3个Tab无数据'],
]
for i, row_data in enumerate(d3):
    for j, val in enumerate(row_data):
        t3.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# ===== 三、需补充数据清单 =====
doc.add_heading('三、需补充的数据（逐项确认）', level=1)

t4 = doc.add_table(rows=11, cols=5, style='Table Grid')
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
h4 = ['#', '数据类别', '涉及页面', '现状', '需确认']
for i, h in enumerate(h4):
    cell = t4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d4 = [
    ['1', '电路拓扑图', '全部6页', '全部是占位图 anquanqinang.png', '每个子应用需要真实拓扑图？还是统一用SVG方块图（data/topology.js已有框架）？还是从应用手册截图？'],
    ['2', '子应用描述(description)', '全部6页', '全部为空字符串 ""', '每个子应用一句话简介，说明该子应用的功能和在系统中的位置'],
    ['3', '推荐系列(series)', '全部6页', '全部为空数组 []', '每个子应用推荐哪些产品系列？（数据已有但仅在SPECS中，需同步到appData）'],
    ['4', '汽车电子 t3-t8 规格', '汽车电子', '车灯/智能驾驶/多媒体/辅助/OBC/充电桩6个Tab完全无SPECS', '是否从对应应用手册提取？还是暂用默认4模块卡片？'],
    ['5', '汽车电子名称对齐', '汽车电子', '"直流转直流-DCDC"与"DC-DC转换器"不匹配', '统一命名规范：appData子应用名 = SPECS键名'],
    ['6', '消费类电子名称对齐', '消费类电子', '"快充输入端专用"与"PD快充-高压输入端"不匹配', '统一命名，决定以哪边为准'],
    ['7', '消费类电子缺失Tab', '消费类电子', '电子笔/数码3C/安防3个Tab无SPECS', '是否需要补充规格数据？'],
    ['8', '下载按钮PDF链接', '全部6页', '按钮存在但链接为空', '对应的选型指南/应用白皮书PDF文件路径'],
    ['9', '技术文章链接', '全部6页', '标题正确但链接为 #', '文章实际URL（官网新闻页 or 公众号链接）'],
    ['10', '设计工具链接', '全部6页', '寿命推算/SPICE/3D-CAD链接为 #', '工具页面URL'],
]
for i, row_data in enumerate(d4):
    for j, val in enumerate(row_data):
        t4.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# ===== 四、汽车电子名称对齐明细 =====
doc.add_heading('四、汽车电子 appData 与 SPECS 名称对齐明细', level=1)

doc.add_paragraph('以下子应用名在 appData 和 SPECS 数据库中不一致，需逐项确认以哪个为准：', style='Normal')

t5 = doc.add_table(rows=11, cols=4, style='Table Grid')
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
h5 = ['#', '所属Tab', 'appData 子应用名（页面显示）', 'SPECS 数据库键名']
for i, h in enumerate(h5):
    cell = t5.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d5 = [
    ['1', '电机驱动', '直流转直流-DCDC', 'DC-DC转换器'],
    ['2', '电机驱动', '电池管理系统-BMS', 'BMS电池管理'],
    ['3', '电机驱动', '高压分线盒-PDU', '（无对应SPECS键）'],
    ['4', '电机驱动', '整车控制-VCU', '（无对应SPECS键）'],
    ['5', '安全部件', '车身稳定系统-ESC', '（无对应SPECS键）'],
    ['6', '安全部件', '座椅通风', '（无对应SPECS键）'],
    ['7', '热管理部件', 'BMS电池管理', '（无对应SPECS键 — 与电机驱动的BMS不同）'],
    ['8', '热管理部件', '冷却风扇控制器', '（无对应SPECS键）'],
    ['9', '热管理部件', '电子水阀', '（无对应SPECS键）'],
    ['10', '热管理部件', '电子水泵', '（无对应SPECS键）'],
]
for i, row_data in enumerate(d5):
    for j, val in enumerate(row_data):
        t5.rows[i+1].cells[j].text = val

doc.add_paragraph('')
doc.add_paragraph('另：t3(新能源车车灯)/t4(智能驾驶8个子应用)/t5(多媒体系统4个子应用)/t6(辅助2个子应用)/t7(车载OBC)/t8(充电桩) 共6个Tab下所有子应用在SPECS中均无对应键，需确认是否补充。', style='Normal')

doc.add_paragraph('')

# ===== 五、会议建议流程 =====
doc.add_heading('五、建议会议流程', level=1)

flow = [
    '1. 逐页过一遍样式 — 先看AI服务器页（数据最完整，7个子应用全部有料号+模块卡片）→ 储能 → 电机驱动 → 仪器仪表 → 汽车电子 → 消费类电子',
    '2. 确认电路拓扑图方案：SVG方块图 / 真实电路截图 / PDF应用手册截图',
    '3. 逐页确认"推荐系列"数据补充来源（哪本应用手册、哪一页）',
    '4. 确认汽车电子缺失的25个子应用是否需要补充SPECS数据，还是用默认模块卡片即可',
    '5. 确认消费类电子子应用命名统一（以产品推广部最终命名为准）',
    '6. 确认下载按钮PDF、技术文章、设计工具的链接地址',
    '7. 明确各方数据交付时间节点',
]

for item in flow:
    doc.add_paragraph(item, style='List Number')

doc.add_paragraph('')
doc.add_paragraph('')

# ===== 附录：6页数据总览 =====
doc.add_heading('附录：6个应用领域页面数据总览', level=1)

t6 = doc.add_table(rows=7, cols=8, style='Table Grid')
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
h6 = ['页面', '一级Tab', '子应用数', 'SPECS覆盖', '模块卡片', '拓扑图', '系列数据', '描述']
for i, h in enumerate(h6):
    cell = t6.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d6 = [
    ['汽车电子', '9个', '36个', '~11个(30%)', '部分有', '占位图', '全空 []', '全空 ""'],
    ['AI服务器', '3个', '7个', '7个(100%)', '全部有', '占位图', '全空 []', '全空 ""'],
    ['消费类电子', '5个', '6个', '0个(0%)', '部分有', '占位图', '全空 []', '全空 ""'],
    ['储能', '3个', '4个', '4个(100%)', '全部有', '占位图', '全空 []', '全空 ""'],
    ['电机驱动', '6个', '6个', '6个(100%)', '全部有', '占位图', '全空 []', '全空 ""'],
    ['仪器仪表', '3个', '9个', '9个(100%)', '全部有', '占位图', '全空 []', '全空 ""'],
]
for i, row_data in enumerate(d6):
    for j, val in enumerate(row_data):
        t6.rows[i+1].cells[j].text = val

doc.add_paragraph('')
doc.add_paragraph('注：所有页面的 description/series/topology 均为占位数据。真实数据仅在 data/system-modules.js 的 SPECS 和 DB 对象中。', style='Normal')

# 保存
output_path = '应用中心内容确认清单.docx'
doc.save(output_path)
print(f'已生成：{output_path}')
