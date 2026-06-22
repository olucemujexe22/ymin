# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

title = doc.add_heading('〇、应用总览页样式确认（application-center.html）', level=1)

doc.add_paragraph('总览页为6个应用领域的入口页面，结构如下：', style='Normal')
doc.add_paragraph('')

# 页面结构表
doc.add_heading('页面结构', level=2)

t1 = doc.add_table(rows=7, cols=3, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['区域', '内容', '确认点']):
    cell = t1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d1 = [
    ['面包屑', '首页 > 应用中心', ''],
    ['标题区', 'h1「应用中心」+ 一段引言描述', '引言文案是否OK？'],
    ['6大应用领域卡片', '3列网格(grid-cols-3)，每卡片含：Material图标 + 领域名 + 简介 + 特性标签 + 推荐系列', '每张卡片的简介、标签、系列名称是否需要修改？详见下方核对表。'],
    ['热门终端应用', '4个快捷入口卡片：安全气囊ECU / ABS防抱死 / ADAS自动驾驶 / GPU加速卡VRM', '① 安全气囊链接到 application-airbag.html（该文件已删除！需移除或替换）\n② ABS/ADAS 链接带 ?tab= 参数跳转\n③ 入口数量和内容是否合适？'],
    ['技术资源横幅', '深蓝底白色字 CTA —「AI服务器PDN优化指南」下载白皮书', '① 只展示AI服务器白皮书是否合适？\n② 下载链接待补充'],
    ['设计工具入口', '3个横排卡片：寿命推算 / SPICE模型 / 3D-CAD', '链接均为 # 占位，需补充实际URL'],
]
for i, row_data in enumerate(d1):
    for j, val in enumerate(row_data):
        t1.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# 6大领域卡片核对表
doc.add_heading('6大领域卡片内容核对', level=2)

doc.add_paragraph('以下内容均为总览页硬编码，修改详情页不会自动同步。需逐项确认。', style='Normal')

t2 = doc.add_table(rows=7, cols=4, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['领域', '当前简介', '当前标签', '当前系列']):
    cell = t2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

d2 = [
    ['汽车电子', '主驱/电驱电控、充电系统、安全部件、ADAS、热管理、智能座舱、车灯、充电桩 —— AEC-Q200合规，耐高温135℃，抗振动。', 'AEC-Q200 / 135℃耐高温 / 极低漏电流', '8个应用子领域 · 10+产品系列'],
    ['AI服务器', 'GPU 加速卡、CPU 供电、PDN 总线架构的大电流、超低 ESR 电容方案，支持 400V/800V DC-Link。', '大电流 / 超低ESR / 小型化', 'VHT · NPX · CW3H系列'],
    ['仪器仪表', '精密测量仪器、工业自动化仪表、医疗检测设备等对精度和稳定性要求极高的电容方案。', '高精度 / 低漏电流 / 长寿命', 'VKM · VPT · LK系列'],
    ['新型电机驱动', '变频器、伺服驱动器、步进电机驱动等工业电机控制系统的 DC-Link 与滤波电容方案。', 'DC-Link / 滤波 / 高纹波', 'VPG · VPX · CW3H系列'],
    ['储能', '光伏逆变器、储能系统、充电桩等高压大容量电容方案，支持高纹波电流与长寿命需求。', '高压大容量 / 高纹波 / 长寿命', 'CW3H · VPG系列'],
    ['消费类电子', '笔记本电脑、智能家居、PD快充、LED照明等消费电子产品的薄型化小型化电容方案。', '薄型化 / 小型化 / 低ESR', 'NPX · VKM · VPT系列'],
]
for i, row_data in enumerate(d2):
    for j, val in enumerate(row_data):
        t2.rows[i+1].cells[j].text = val

doc.add_paragraph('')

# 需修复问题
doc.add_heading('⚠️ 总览页需修复的问题', level=2)

issues = [
    '1. application-airbag.html 已删除 —「热门终端应用」中"安全气囊ECU"卡片链接到不存在的页面，需移除该卡片或替换为有效链接。',
    '2. ?tab= 参数跳转 — ABS 卡片链接到 application-automotive.html?tab=safety，但汽车电子页只认 t0~t8 键名（safety 无效），已临时改为 ?tab=t1。需最终确认。',
    '3. 同上 — ADAS 卡片链接 ?tab=adas 无效，已临时改为 ?tab=t4。需最终确认。',
    '4. 白皮书横幅 — 下载按钮无实际链接，需提供PDF文件地址。',
    '5. 设计工具3个入口 — 链接均为 #，需补充寿命推算/SPICE/3D-CAD 的实际URL。',
    '6. 6张领域卡片的数据是硬编码在HTML中 — 如果后续修改了领域简介/标签/系列名，需同步修改总览页HTML。建议后续改为从共享数据源渲染。',
]
for item in issues:
    doc.add_paragraph(item, style='Normal')

output_path = '应用总览页确认清单.docx'
doc.save(output_path)
print(f'已生成：{output_path}')
