# -*- coding: utf-8 -*-
"""读取应用中心内容确认清单.docx，提取结构信息"""
from docx import Document

doc = Document('应用中心内容确认清单.docx')

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'Normal'
        print(f"[P{i}|{style}] {para.text[:120]}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.replace('\n', ' | ')[:60] for cell in row.cells]
        print(f"  R{ri}: {' || '.join(cells)}")
